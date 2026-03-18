import os
from io import BytesIO
from pypdf import PdfReader
from docx import Document


def read_pdf(file_path):
    text_parts = []
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    except Exception:
        return ""
    return "\n".join(text_parts)


def read_pdf_bytes(file_bytes):
    text_parts = []
    try:
        reader = PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    except Exception:
        return ""
    return "\n".join(text_parts)


def read_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        return ""


def read_docx_bytes(file_bytes):
    try:
        doc = Document(BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        return ""


def read_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read()
    except Exception:
        return ""


def read_txt_bytes(file_bytes):
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")
    except Exception:
        return ""


def read_document(file_path):
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return read_pdf(file_path)
    if extension == ".docx":
        return read_docx(file_path)
    if extension == ".txt":
        return read_txt(file_path)

    return ""


def read_document_bytes(file_bytes, filename):
    extension = os.path.splitext(filename)[1].lower()

    if extension == ".pdf":
        return read_pdf_bytes(file_bytes)
    if extension == ".docx":
        return read_docx_bytes(file_bytes)
    if extension == ".txt":
        return read_txt_bytes(file_bytes)

    return ""


def split_into_fragments(text, fragment_size=1200):
    clean_text = " ".join(text.split())
    if not clean_text:
        return []

    fragments = []
    start = 0
    while start < len(clean_text):
        end = start + fragment_size
        fragments.append(clean_text[start:end])
        start = end

    return fragments